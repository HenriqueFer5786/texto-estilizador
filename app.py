from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
from bs4 import BeautifulSoup

app = Flask(__name__)
CORS(app)

# Variável global para armazenar o texto estilizado
global_stylized_text = ""

@app.route('/generate-docx', methods=['POST'])
def generate_docx():
    try:
        global global_stylized_text
        
        # Verifica se há texto estilizado armazenado
        if not global_stylized_text:
            return jsonify({'error': 'Nenhum texto estilizado disponível para gerar o documento'}), 400
        
        # Criando um documento Word
        document = Document()
        document.add_heading('Texto Estilizado', level=1)
        
        # Utilizando BeautifulSoup para analisar o HTML e aplicar estilos do Word
        soup = BeautifulSoup(global_stylized_text, 'html.parser')

        # Percorrendo os elementos do HTML e convertendo para estilos do Word
        for element in soup.find_all(recursive=False):
            if element.name == 'p':
                # Criando um novo parágrafo para cada tag <p>
                paragraph = document.add_paragraph()
                for sub_element in element.contents:
                    run = paragraph.add_run(sub_element.text if sub_element.string else sub_element)
                    if sub_element.name == 'strong' or sub_element.name == 'b':
                        run.bold = True
                    if sub_element.name == 'em' or sub_element.name == 'i':
                        run.italic = True
            elif element.name == 'h1':
                # Adicionando um título de nível 1
                document.add_heading(element.text, level=1)
            elif element.name == 'h2':
                # Adicionando um título de nível 2
                document.add_heading(element.text, level=2)
            elif element.name == 'ul':
                # Adicionando uma lista não ordenada
                for li in element.find_all('li'):
                    document.add_paragraph(li.text, style='ListBullet')
            elif element.name == 'ol':
                # Adicionando uma lista ordenada
                for li in element.find_all('li'):
                    document.add_paragraph(li.text, style='ListNumber')
            elif element.name == 'table':
                # Adicionando uma tabela
                rows = element.find_all('tr')
                if rows:
                    # Criando uma tabela com a quantidade de linhas e colunas encontradas
                    num_cols = len(rows[0].find_all(['td', 'th']))
                    table = document.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    for row_idx, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for col_idx, cell in enumerate(cells):
                            cell_text = cell.get_text(strip=True)
                            cell_paragraph = table.cell(row_idx, col_idx).paragraphs[0]
                            run = cell_paragraph.add_run(cell_text)
                            
                            # Aplicar negrito no cabeçalho da tabela
                            if cell.name == 'th':
                                run.bold = True
            elif element.name == 'strong' or element.name == 'b':
                # Negrito fora de um parágrafo
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.text)
                run.bold = True
            elif element.name == 'em' or element.name == 'i':
                # Itálico fora de um parágrafo
                paragraph = document.add_paragraph()
                run = paragraph.add_run(element.text)
                run.italic = True
            elif element.name is None and element.string:
                # Texto puro fora de uma tag específica
                document.add_paragraph(element.string)

        # Salvando o documento em um buffer de memória
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # Enviando o arquivo .docx como resposta
        return send_file(buffer, as_attachment=True, download_name='texto_estilizado.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        print(f"Erro ao gerar documento: {e}")
        return jsonify({'error': 'Erro ao gerar o documento'}), 500

@app.route('/stylize-text', methods=['POST'])
def stylize_text():
    try:
        global global_stylized_text
        
        # Recebendo o texto do cliente
        data = request.get_json()
        markdown_text = data.get('text', '')
        print(f"Texto recebido para estilização: {markdown_text}")
        
        # Convertendo o texto markdown para HTML
        global_stylized_text = markdown.markdown(markdown_text, extensions=['tables'])
        print(f"Texto estilizado (HTML): {global_stylized_text}")

        # Retornando o texto estilizado como JSON
        return jsonify({'stylized_text': global_stylized_text})
    except Exception as e:
        print(f"Erro ao estilizar texto: {e}")
        return jsonify({'error': 'Erro ao estilizar o texto'}), 500

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=int(os.environ.get("PORT", 5000)))
